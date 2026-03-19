from flask import Flask, render_template, request, send_file, jsonify, redirect, flash, session, url_for
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.pdfbase.pdfmetrics import stringWidth
from document_scanner import extract_text_from_file, simple_extract_data
import io
import os
import re
import zipfile
from werkzeug.utils import secure_filename

from werkzeug.utils import secure_filename
try:
    from pypdf import PdfReader, PdfWriter
except Exception:
    try:
        from PyPDF2 import PdfReader, PdfWriter
    except Exception:
        PdfReader = None
        PdfWriter = None

app = Flask(__name__)
UPLOAD_FOLDER = os.path.join(app.root_path, "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.secret_key = "ponce_secret_key"
USERNAME = "admin"
PASSWORD = "ponce"


def clean_text(value):
    return (value or "").strip()


def clean_facts_of_loss(text):
    text = clean_text(text)
    replacements = {
        "client was client was": "client was",
        "our client was client was": "our client was",
        "was was": "was",
        "the the": "the",
        "rear-ended from the rear": "rear-ended",
        "collision collision": "collision",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
        text = text.replace(bad.title(), good.title())
    return " ".join(text.split())


def parse_money(value):
    value = clean_text(value).replace("$", "").replace(",", "")
    if not value:
        return 0.0
    try:
        return float(value)
    except ValueError:
        return 0.0


def money(value):
    return "${:,.2f}".format(value or 0.0)

def unique_join_paragraphs(existing, new_text):
    existing = clean_text(existing)
    new_text = clean_text(new_text)

    if not new_text:
        return existing

    if not existing:
        return new_text

    existing_parts = [p.strip() for p in existing.split("\n\n") if p.strip()]
    new_parts = [p.strip() for p in new_text.split("\n\n") if p.strip()]

    seen = {p.lower() for p in existing_parts}
    for part in new_parts:
        if part.lower() not in seen:
            existing_parts.append(part)
            seen.add(part.lower())

    return "\n\n".join(existing_parts)


def dedupe_provider_rows(providers):
    cleaned = []
    seen = set()

    for p in providers:
        name = clean_text(p.get("name") or p.get("provider_name"))
        amount = clean_text(str(p.get("amount", "")))

        if not name and not amount:
            continue

        key = re.sub(r"\s+", " ", name).strip().lower()
        if key and key in seen:
            continue

        if key:
            seen.add(key)

        cleaned.append({
            "name": name,
            "amount": amount
        })

    return cleaned


def safe_filename(value):
    value = clean_text(value)
    if not value:
        return "Demand_Letter"
    value = re.sub(r"[^A-Za-z0-9._ -]+", "", value)
    value = value.replace(" ", "_")
    return value or "Demand_Letter"


def add_paragraph(document, text="", bold=False, font_size=11, space_after=8):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def find_existing_path(candidates):
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def get_letterhead_pdf_path():
    return find_existing_path([
        "TPL Letterhead.pdf",
        "letterhead.pdf",
        "static/branding/TPL Letterhead.pdf",
        "static/branding/letterhead.pdf",
        "static/branding/Letterhead.pdf",
    ])


def get_logo_path():
    return find_existing_path([
        "static/branding/logo.png",
        "static/branding/logo.jpg",
        "static/branding/logo.jpeg",
    ])


def build_letter_data(form):
    recipient_name = clean_text(form.get("recipient_name"))
    adjuster_name = clean_text(form.get("adjuster_name"))
    client_name = clean_text(form.get("client_name"))
    claim_number = clean_text(form.get("claim_number"))
    loss_date = clean_text(form.get("loss_date"))
    deadline = clean_text(form.get("deadline"))

    facts_of_loss = clean_facts_of_loss(form.get("facts_of_loss"))
    treatment_text = clean_text(form.get("treatment_summary"))
    non_economic_text = clean_text(form.get("non_economic_damages"))
    damages_explanation = clean_text(form.get("damages_explanation"))

    provider_names = form.getlist("provider_name[]")
    provider_amounts = form.getlist("provider_amount[]")

    raw_providers = []
    for name, amount in zip(provider_names, provider_amounts):
        clean_name = clean_text(name)
        clean_amount_raw = clean_text(amount)

        if not clean_name and not clean_amount_raw:
            continue

        raw_providers.append({
            "name": clean_name,
            "amount": clean_amount_raw
        })

    deduped_rows = dedupe_provider_rows(raw_providers)

    providers = []
    for row in deduped_rows:
        clean_name = clean_text(row.get("name"))
        clean_amount_raw = clean_text(row.get("amount"))

        try:
            clean_amount = float(clean_amount_raw.replace(",", "").replace("$", "")) if clean_amount_raw else 0.0
        except ValueError:
            clean_amount = 0.0

        providers.append({
            "provider_name": clean_name or "Provider",
            "amount": clean_amount,
            "amount_display": money(clean_amount)
        })

    medical_expenses = parse_money(form.get("medical_expenses"))
    lost_wages = parse_money(form.get("lost_wages"))

    provider_total = sum(item["amount"] for item in providers)
    if provider_total > 0:
        medical_expenses = provider_total

    multiplier_value = clean_text(form.get("multiplier"))
    try:
        multiplier_num = float(multiplier_value) if multiplier_value else 3.0
    except ValueError:
        multiplier_num = 3.0

    economic_total = medical_expenses + lost_wages
    suggested_amount = economic_total * multiplier_num

    policy_limit_checked = form.get("policy_limit_demand") == "yes"

    if facts_of_loss:
        liability_text = (
            "Based on the facts described above, liability for this collision rests with your insured. "
            "The at-fault driver failed to exercise ordinary care in the operation of the vehicle, and "
            "that negligence was the direct and proximate cause of our client's injuries and damages."
        )
    else:
        liability_text = (
            "Liability is clear. The available facts show that the collision was caused by the negligence "
            "of the at-fault driver, who failed to operate the vehicle in a reasonably safe manner under "
            "the circumstances."
        )

    if policy_limit_checked:
        demand_type = "Policy Limits Demand"
        suggested_settlement_display = "Policy Limits"
        conclusion_text = (
            "Based on the clear liability of your insured, the nature and extent of our client's injuries, "
            "the medical treatment required, and the resulting damages, we hereby demand tender of all "
            "available policy limits within the time allowed."
        )
    else:
        demand_type = "Settlement Demand"
        suggested_settlement_display = money(suggested_amount)
        conclusion_text = (
            "Based on the clear liability of your insured, the nature and severity of this collision, "
            "the medical treatment incurred, the economic damages sustained, and the ongoing impact on "
            f"our client's quality of life, we hereby demand {suggested_settlement_display} in full "
            "and final settlement of this claim."
        )

    return {
        "recipient_name": recipient_name,
        "adjuster_name": adjuster_name,
        "client_name": client_name,
        "claim_number": claim_number,
        "loss_date": loss_date,
        "deadline": deadline,
        "facts_of_loss": facts_of_loss,
        "liability_text": liability_text,
        "treatment_text": treatment_text,
        "non_economic_text": non_economic_text,
        "damages_explanation": damages_explanation,
        "providers": providers,
        "medical_expenses": medical_expenses,
        "lost_wages": lost_wages,
        "multiplier_num": multiplier_num,
        "economic_total": economic_total,
        "suggested_settlement_display": suggested_settlement_display,
        "policy_limit_checked": policy_limit_checked,
        "demand_type": demand_type,
        "conclusion_text": conclusion_text,
    }


def build_docx(letter_data):
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    add_paragraph(
        document,
        "This letter is written for the sole purpose of settlement and is not intended for use for any other purpose without the express written consent of this office. Accordingly, this letter and the material included herein shall not, to the extent allowed by law, be admitted as evidence.",
        font_size=11,
        space_after=12
    )

    recipient_line = letter_data["recipient_name"] or letter_data["adjuster_name"] or "Adjuster"
    add_paragraph(document, f"Dear {recipient_line}:", font_size=11, space_after=12)

    intro = (
        f"Texas Ponce Law, PLLC has been retained to represent {letter_data['client_name'] or 'our client'} "
        f"for injuries sustained as a result of a motor vehicle collision that occurred on "
        f"{letter_data['loss_date'] or 'the date of loss'}."
    )
    add_paragraph(document, intro, font_size=11, space_after=12)

    add_paragraph(
        document,
        "We have been instructed to recover compensation for the debt owed and personal injuries sustained as a direct result of this collision.",
        font_size=11,
        space_after=12
    )

    if letter_data["claim_number"]:
        add_paragraph(document, f"Claim Number: {letter_data['claim_number']}", font_size=11, space_after=12)

    if letter_data["deadline"]:
        add_paragraph(document, f"This settlement offer is open until {letter_data['deadline']}.", font_size=11, space_after=12)

    add_paragraph(
        document,
        "Enclosed for your review are my client's medical records, bills, and supporting documentation.",
        font_size=11,
        space_after=12
    )

    add_paragraph(document, "FACTS OF LOSS", bold=True, font_size=18, space_after=8)
    add_paragraph(document, letter_data["facts_of_loss"] or "N/A", font_size=11, space_after=16)

    add_paragraph(document, "LIABILITY", bold=True, font_size=18, space_after=8)
    add_paragraph(document, letter_data["liability_text"], font_size=11, space_after=16)

    add_paragraph(document, "TREATMENT AND INJURIES", bold=True, font_size=18, space_after=8)
    treatment_body = letter_data["treatment_text"] or "N/A"
    add_paragraph(document, treatment_body, font_size=11, space_after=16)

    add_paragraph(document, "PAST MEDICAL EXPENSES", bold=True, font_size=18, space_after=8)
    if letter_data["providers"]:
        for provider in letter_data["providers"]:
            add_paragraph(
                document,
                f"{provider['provider_name']}: {provider['amount_display']}",
                font_size=11,
                space_after=4
            )
    else:
        add_paragraph(document, "No provider bills listed.", font_size=11, space_after=8)

    add_paragraph(document, "", font_size=11, space_after=2)
    add_paragraph(document, f"TOTAL MEDICAL EXPENSES: {money(letter_data['medical_expenses'])}", bold=True, font_size=11, space_after=16)

    add_paragraph(document, "NON-ECONOMIC DAMAGES", bold=True, font_size=18, space_after=8)
    non_econ_body = (
        letter_data["non_economic_text"]
        or letter_data["damages_explanation"]
        or "As a result of the collision and resulting injuries, our client has endured physical pain, emotional distress, disruption to daily life, and ongoing limitations affecting overall quality of life."
    )
    add_paragraph(document, non_econ_body, font_size=11, space_after=16)

    add_paragraph(document, "CONCLUSION", bold=True, font_size=18, space_after=8)
    if not letter_data["policy_limit_checked"]:
        add_paragraph(document, "Based on:", font_size=11, space_after=6)
        add_paragraph(document, "• the clear liability of your insured", font_size=11, space_after=2)
        add_paragraph(document, "• the severity of the collision", font_size=11, space_after=2)
        add_paragraph(document, "• the objective diagnostic findings", font_size=11, space_after=2)
        add_paragraph(document, "• the economic damages sustained", font_size=11, space_after=2)
        add_paragraph(document, "• the ongoing impact on our client's quality of life", font_size=11, space_after=12)

    add_paragraph(document, letter_data["conclusion_text"], font_size=11, space_after=12)

    if letter_data["deadline"]:
        add_paragraph(document, f"This offer will remain open until {letter_data['deadline']}.", font_size=11, space_after=12)

    add_paragraph(document, "Please contact our office if you have any questions.", font_size=11, space_after=12)
    add_paragraph(document, "Sincerely,", font_size=11, space_after=18)
    add_paragraph(document, "Texas Ponce Law, PLLC", font_size=11, space_after=0)

    return document

def build_content_pdf(letter_data):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=LETTER)
    page_width, page_height = LETTER

    left_margin = 95
    right_margin = 95
    first_page_top_margin = 195
    other_page_top_margin = 72
    bottom_margin = 70
    usable_width = page_width - left_margin - right_margin
    y = page_height - first_page_top_margin
   
    c.setFillColorRGB(1, 1, 1)
    c.rect(45, 20, 240, 75, fill=1, stroke=0)
    c.setFillColorRGB(0, 0, 0)

    def new_page():
        nonlocal y
        c.showPage()
        y = page_height - other_page_top_margin

    def ensure_space(amount):
        nonlocal y
        if y - amount < bottom_margin:
            new_page()

    def wrap_text(text, font_name, font_size):
        words = str(text or "").split()
        if not words:
            return []

        lines = []
        current = ""

        for word in words:
            test = word if not current else f"{current} {word}"
            if stringWidth(test, font_name, font_size) <= usable_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = word

        if current:
            lines.append(current)

        return lines

    def draw_paragraph(text, font_name="Times-Roman", font_size=11, leading=15, space_after=8):
        nonlocal y
        text = str(text or "").strip()
        if not text:
            y -= space_after
            return

        lines = wrap_text(text, font_name, font_size)
        ensure_space((len(lines) * leading) + space_after)

        c.setFont(font_name, font_size)
        for line in lines:
            c.drawString(left_margin, y, line)
            y -= leading

        y -= space_after

    def draw_heading(text, font_size=15):
        nonlocal y
        ensure_space(34)
        c.setFont("Times-Bold", font_size)
        c.drawString(left_margin, y, text)
        y -= 18
        c.setLineWidth(0.3)
        c.line(left_margin, y, page_width - right_margin, y)
        y -= 12
       
    def draw_provider_table(providers):
        nonlocal y
        col1 = left_margin
        col2 = left_margin + 285

        rows = providers[:] if providers else [{"provider_name": "No provider bills listed.", "amount_display": ""}]
        row_height = 24
        total_needed = 34 + (len(rows) * row_height) + 24
        ensure_space(total_needed)

        c.setFont("Times-Bold", 13)
        c.drawString(col1, y, "Provider")
        c.drawString(col2, y, "Amount")
        y -= 22

        c.setFont("Times-Roman", 13)
        for provider in rows:
            c.drawString(col1, y, provider.get("provider_name", ""))
            c.drawString(col2, y, provider.get("amount_display", ""))
            y -= row_height

    draw_paragraph(
        "This letter is written for the sole purpose of settlement and is not intended for use for any other purpose without the express written consent of this office. Accordingly, this letter and the material included herein shall not, to the extent allowed by law, be admitted as evidence.",
        font_size=12,
        leading=17,
        space_after=18
    )

    recipient_line = letter_data["recipient_name"] or letter_data["adjuster_name"] or "Adjuster"
    draw_paragraph(f"Dear {recipient_line}:", font_size=12, leading=17, space_after=18)

    intro = (
        f"Texas Ponce Law, PLLC has been retained to represent {letter_data['client_name'] or 'our client'} "
        f"for injuries sustained as a result of a motor vehicle collision that occurred on "
        f"{letter_data['loss_date'] or 'the date of loss'}."
    )
    draw_paragraph(intro, font_size=12, leading=17, space_after=18)

    draw_paragraph(
        "We have been instructed to recover compensation for the debt owed and personal injuries sustained as a direct result of this collision.",
        font_size=12,
        leading=17,
        space_after=18
    )

    if letter_data["deadline"]:
        draw_paragraph(
            f"This settlement offer is open until {letter_data['deadline']}.",
            font_size=12,
            leading=17,
            space_after=18
        )

    draw_paragraph(
        "Enclosed for your review are my client's medical records, bills, and supporting documentation.",
        font_size=12,
        leading=17,
        space_after=18
    )

    draw_heading("FACTS OF LOSS")
    draw_paragraph(letter_data["facts_of_loss"] or "N/A", font_size=13, leading=19, space_after=22)

    draw_heading("LIABILITY")
    draw_paragraph(letter_data["liability_text"], font_size=13, leading=19, space_after=22)

    draw_heading("TREATMENT AND INJURIES")
    treatment_body = letter_data["treatment_text"] or "N/A"
    draw_paragraph(treatment_body, font_size=13, leading=19, space_after=22)

    draw_heading("PAST MEDICAL EXPENSES")
    draw_paragraph("Below is a summary of the medical bills incurred as a result of this collision.", font_size=13, leading=19, space_after=14)
    draw_provider_table(letter_data["providers"])
    y -= 12
    draw_paragraph(f"TOTAL MEDICAL EXPENSES: {money(letter_data['medical_expenses'])}", font_name="Times-Bold", font_size=13, leading=19, space_after=22)

    draw_heading("NON-ECONOMIC DAMAGES")
    non_econ_body = (
        letter_data["non_economic_text"]
        or letter_data["damages_explanation"]
        or "As a result of the collision and resulting injuries, our client has endured physical pain, emotional distress, disruption to daily life, and ongoing limitations affecting overall quality of life."
    )
    draw_paragraph(non_econ_body, font_size=11, leading=15, space_after=18)

    draw_heading("CONCLUSION")
    if not letter_data["policy_limit_checked"]:
        draw_paragraph("Based on:", font_size=11, leading=15, space_after=6)
        draw_paragraph("• the clear liability of your insured", font_size=11, leading=15, space_after=2)
        draw_paragraph("• the severity of the collision", font_size=11, leading=15, space_after=2)
        draw_paragraph("• the objective diagnostic findings", font_size=11, leading=15, space_after=2)
        draw_paragraph("• the economic damages sustained", font_size=11, leading=15, space_after=2)
        draw_paragraph("• the ongoing impact on our client's quality of life", font_size=11, leading=15, space_after=12)

    draw_paragraph(letter_data["conclusion_text"], font_name="Times-Bold", font_size=11, leading=15, space_after=14)

    if letter_data["deadline"]:
        draw_paragraph(
            f"This offer will remain open until {letter_data['deadline']}.",
            font_size=11,
            leading=15,
            space_after=14
        )

    if y < 90:
        new_page()

    c.setFont("Times-Roman", 9)
    c.drawString(left_margin, 58, "Mailing Address:")
    c.drawString(left_margin, 46, "5900 Balcones Drive, #16604")
    c.drawString(left_margin, 34, "Austin, TX 78731")

    c.save()
    buffer.seek(0)
    return buffer

def apply_letterhead_overlay(content_pdf_bytes):
    letterhead_path = get_letterhead_pdf_path()

    if not letterhead_path or PdfReader is None or PdfWriter is None:
        return content_pdf_bytes

    try:
        base_reader = PdfReader(letterhead_path)
        content_reader = PdfReader(io.BytesIO(content_pdf_bytes))
        writer = PdfWriter()

        if len(base_reader.pages) == 0:
            return content_pdf_bytes

        first_template = base_reader.pages[0]

        for index, content_page in enumerate(content_reader.pages):
            if index == 0:
                new_page = writer.add_blank_page(
                    width=first_template.mediabox.width,
                    height=first_template.mediabox.height
                )
                new_page.merge_page(first_template)
                new_page.merge_page(content_page)
            else:
                writer.add_page(content_page)

        output = io.BytesIO()
        writer.write(output)
        return output.getvalue()
    except Exception as e:
        print("PDF overlay error:", e)
        return content_pdf_bytes

def build_pdf_bytes(letter_data):
    content_buffer = build_content_pdf(letter_data)
    raw_pdf_bytes = content_buffer.getvalue()
    final_pdf_bytes = apply_letterhead_overlay(raw_pdf_bytes)
    return final_pdf_bytes


@app.route("/login", methods=["GET", "POST"])
def login():
    error = None

    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")

        if username == USERNAME and password == PASSWORD:
            session["logged_in"] = True
            return redirect(url_for("index"))
        else:
            error = "Invalid login credentials"

    return render_template("login.html", error=error)


@app.route("/", methods=["GET"])
def index():
    if not session.get("logged_in"):
        return redirect(url_for("login"))

    scanned_data = session.pop("scanned_data", None)
    return render_template("form.html", scanned_data=scanned_data)


@app.route("/generate", methods=["POST"])
def generate():
    if not session.get("logged_in"):
        return redirect(url_for("login"))

    letter_data = build_letter_data(request.form)

    document = build_docx(letter_data)
    docx_stream = io.BytesIO()
    document.save(docx_stream)
    docx_stream.seek(0)

    pdf_bytes = build_pdf_bytes(letter_data)

    demand_label = "Policy_Limits_Demand" if letter_data["policy_limit_checked"] else "Settlement_Demand"
    client_part = safe_filename(letter_data["client_name"] or "Client")

    docx_filename = f"{client_part}_{demand_label}.docx"
    pdf_filename = f"{client_part}_{demand_label}.pdf"
    zip_filename = f"{client_part}_{demand_label}_Bundle.zip"

    zip_stream = io.BytesIO()
    with zipfile.ZipFile(zip_stream, "w", zipfile.ZIP_DEFLATED) as zf:
        folder_name = "Texas_Ponce_Demand/"

        zf.writestr(folder_name + docx_filename, docx_stream.getvalue())
        zf.writestr(folder_name + pdf_filename, pdf_bytes)

        icon_path = "static/branding/icon.icns"
        if os.path.exists(icon_path):
            with open(icon_path, "rb") as icon_file:
                zf.writestr(folder_name + ".VolumeIcon.icns", icon_file.read())

    zip_stream.seek(0)

    return send_file(
        zip_stream,
        as_attachment=True,
        download_name=zip_filename,
        mimetype="application/zip"
    )


@app.route("/preview-data", methods=["POST"])
def preview_data():
    if not session.get("logged_in"):
        return redirect(url_for("login"))

    letter_data = build_letter_data(request.form)
    return jsonify({
        "demand_type": letter_data["demand_type"],
        "medical_expenses": money(letter_data["medical_expenses"]),
        "lost_wages": money(letter_data["lost_wages"]),
        "suggested_settlement": letter_data["suggested_settlement_display"]
    })


@app.route("/scan-documents", methods=["GET", "POST"])
def scan_documents():
    if not session.get("logged_in"):
        return redirect(url_for("login"))

    if request.method == "POST":
        files = request.files.getlist("documents")

        combined_data = {
            "client_name": "",
            "claim_number": "",
            "loss_date": "",
            "facts_of_loss": "",
            "treatment_summary": "",
            "providers": [],
            "review": {
                "client_name": "",
                "claim_number": "",
                "loss_date": "",
                "facts_of_loss": "",
                "treatment_summary": "",
                "providers_text": "",
                "chronology_text": "",
                "objective_findings_text": "",
                "imaging_summary_text": "",
                "high_value_text": "",
                "treatment_gaps": "",
                "demand_treatment_section": "",
                "providers": []
            }
        }

        for file in files:
            if file and file.filename:
                filename = secure_filename(file.filename)
                lower_name = filename.lower()

                if not (lower_name.endswith(".pdf") or lower_name.endswith(".docx")):
                    continue

                filepath = os.path.join(UPLOAD_FOLDER, filename)
                file.save(filepath)

                text = extract_text_from_file(filepath)
                data = simple_extract_data(text)

                if data.get("client_name") and not combined_data["client_name"]:
                    combined_data["client_name"] = data["client_name"]
                    combined_data["review"]["client_name"] = data["review"]["client_name"]

                if data.get("claim_number") and not combined_data["claim_number"]:
                    combined_data["claim_number"] = data["claim_number"]
                    combined_data["review"]["claim_number"] = data["review"]["claim_number"]

                if data.get("loss_date") and not combined_data["loss_date"]:
                    combined_data["loss_date"] = data["loss_date"]
                    combined_data["review"]["loss_date"] = data["review"]["loss_date"]

                if data.get("facts_of_loss"):
                    combined_data["facts_of_loss"] = unique_join_paragraphs(
                        combined_data["facts_of_loss"],
                        data["facts_of_loss"]
                    )
                    combined_data["review"]["facts_of_loss"] = combined_data["facts_of_loss"]

                if data.get("treatment_summary"):
                    combined_data["treatment_summary"] = unique_join_paragraphs(
                        combined_data["treatment_summary"],
                        data["treatment_summary"]
                    )
                    combined_data["review"]["treatment_summary"] = combined_data["treatment_summary"]

                for p in data.get("providers", []):
                    combined_data["providers"].append(p)

                combined_data["providers"] = dedupe_provider_rows(combined_data["providers"])
                combined_data["review"]["providers"] = combined_data["providers"]

                for key in [
                    "providers_text",
                    "chronology_text",
                    "objective_findings_text",
                    "imaging_summary_text",
                    "high_value_text"
                ]:
                    new_text = clean_text(data["review"].get(key, ""))
                    if new_text:
                        combined_data["review"][key] = unique_join_paragraphs(
                            combined_data["review"][key],
                            new_text
                        )

                if data["review"].get("treatment_gaps"):
                    combined_data["review"]["treatment_gaps"] = data["review"]["treatment_gaps"]

        if combined_data["treatment_summary"]:
            combined_data["review"]["demand_treatment_section"] = combined_data["treatment_summary"]

        session["scan_review_data"] = combined_data
        return redirect(url_for("scan_review"))

    return render_template("upload_scan.html")

@app.route("/scan-review", methods=["GET", "POST"])
def scan_review():
    if not session.get("logged_in"):
        return redirect(url_for("login"))

    data = session.get("scan_review_data")

    if not data:
        return redirect(url_for("scan_documents"))

    if request.method == "POST":
        applied = {
            "client_name": request.form.get("client_name", "").strip(),
            "claim_number": request.form.get("claim_number", "").strip(),
            "loss_date": request.form.get("loss_date", "").strip(),
            "facts_of_loss": request.form.get("facts_of_loss", "").strip(),
            "treatment_summary": request.form.get("treatment_summary", "").strip(),
            "providers": []
        }

        provider_names = request.form.getlist("provider_name[]")
        provider_amounts = request.form.getlist("provider_amount[]")

        for name, amount in zip(provider_names, provider_amounts):
            name = name.strip()
            amount = amount.strip()
            if name or amount:
                applied["providers"].append({
                    "name": name,
                    "amount": amount
                })

        session["scanned_data"] = applied
        session.pop("scan_review_data", None)
        return redirect(url_for("index"))

    return render_template("scan_review.html", data=data)


if __name__ == "__main__":
    app.run(debug=True)    
